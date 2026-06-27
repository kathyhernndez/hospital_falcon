"""
Motor de extracción y limpieza de datos desde archivos Excel, Word y PDF.
=========================================================================
Cada función recibe la ruta del archivo y devuelve una tupla:
    (records: list[dict], errors: list[str])

Reglas implementadas:
1. Campos Obligatorios: Nombre, Apellido, Hospital. Si falta alguno, se ignora el registro.
2. Campos Opcionales con validación:
    - Cedula: solo números o prefijos V-/E-
    - Edad: entero (1-120)
    - Contacto: número telefónico
    - Sector: texto libre
"""

import os
import re
from typing import Tuple, List, Dict

import pandas as pd
from docx import Document as DocxDocument
from pypdf import PdfReader


# ---------------------------------------------------------------------------
# Mapeo flexible de columnas para Excel y Tablas de Word
# ---------------------------------------------------------------------------
_COLUMN_MAP = {
    # Nombres combinados (más específicos, se evalúan primero al ordenar por longitud en _map_columns)
    "nombres y apellidos": "nombre_completo",
    "apellidos y nombres": "nombre_completo",
    "nombre y apellido":   "nombre_completo",
    "nombre y apellidos":  "nombre_completo",
    "apellido y nombre":   "nombre_completo",
    "apellidos y nombre":  "nombre_completo",
    "nombre completo":     "nombre_completo",
    "nombre_completo":     "nombre_completo",
    "fullname":            "nombre_completo",
    "full name":           "nombre_completo",

    # Posibles nombres de columna -> campo interno
    "nombre":    "nombre",
    "nombres":   "nombre",
    "name":      "nombre",
    "apellido":  "apellido",
    "apellidos": "apellido",
    "last name": "apellido",
    "cedula":         "cedula",
    "cédula":         "cedula",
    "ci":             "cedula",
    "c.i.":           "cedula",
    "c.i":            "cedula",
    "identidad":      "cedula",
    "edad":    "edad",
    "age":     "edad",
    "sector":       "sector",
    "ciudad":       "sector",
    "ubicacion":    "sector",
    "ubicación":    "sector",
    "direccion":    "sector",
    "dirección":    "sector",
    "hospital":        "hospital",
    "centro":          "hospital",
    "centro de salud": "hospital",
    "centro_salud":    "hospital",
    "contacto":          "contacto",
    "telefono":          "contacto",
    "teléfono":          "contacto",
    "telefono_contacto": "contacto",
    "phone":             "contacto",
}


def _normalize_col(col: str) -> str:
    """Convierte a minúsculas, quita tildes y espacios/guiones."""
    col = str(col).strip().lower()
    replacements = (
        ("á", "a"), ("é", "e"), ("í", "i"), ("ó", "o"), ("ú", "u"),
        ("_", " "), ("-", " ")
    )
    for a, b in replacements:
        col = col.replace(a, b)
    return col


def _map_columns(columns: list) -> dict:
    """Crea un diccionario {columna_original: campo_interno}."""
    mapping = {}
    # Ordenar las claves por longitud descendente para buscar coincidencias más largas primero
    sorted_keys = sorted(_COLUMN_MAP.keys(), key=len, reverse=True)
    
    for col in columns:
        normalized = _normalize_col(col)
        # Búsqueda de coincidencia (substring o exacta)
        for key in sorted_keys:
            if key in normalized:
                mapping[col] = _COLUMN_MAP[key]
                break
    return mapping


# ---------------------------------------------------------------------------
# Expresiones Regulares para Validación de Opcionales
# ---------------------------------------------------------------------------
# Cédula: Empieza opcionalmente por V-, E-, v, e, seguido de 6 a 9 dígitos
_REGEX_CEDULA = re.compile(r"^(?:[VEve][- ]?)?\d{6,9}$")
# Teléfono: +XX y/o código de área, seguido de números
_REGEX_TELEFONO = re.compile(r"^(?:\+?\d{1,3}[- ]?)?\(?\d{3,4}\)?[- ]?\d{3}[- ]?\d{4}$|^\d{10,14}$")


# ---------------------------------------------------------------------------
# Validación Común
# ---------------------------------------------------------------------------
def _validate_record(data: dict, row_num: int) -> Tuple[dict | None, str | None]:
    """
    Aplica las reglas estrictas de campos obligatorios y opcionales.
    Devuelve (registro_limpio, None) si es válido, o (None, error_msg) si no.
    """
    hospital = str(data.get("hospital", "") or "").strip()
    # Clean leading sequence numbers from hospital name (e.g. "1 Hospital Universitario...")
    hospital = re.sub(r"^\d+[-\s]+", "", hospital)

    nombre_completo_raw = str(data.get("nombre_completo", "") or "").strip()
    if nombre_completo_raw:
        nombre_completo = nombre_completo_raw
    else:
        nombre = str(data.get("nombre", "") or "").strip()
        apellido = str(data.get("apellido", "") or "").strip()
        if not nombre:
            return None, f"Fila/Bloque {row_num}: Falta el 'Nombre' (Obligatorio)."
        if not apellido:
            return None, f"Fila/Bloque {row_num}: Falta el 'Apellido' de {nombre} (Obligatorio)."
        nombre_completo = f"{nombre} {apellido}".strip()

    # 1. Campos Obligatorios
    if not hospital:
        return None, f"Fila/Bloque {row_num}: Falta el 'Centro de Salud' para {nombre_completo} (Obligatorio)."

    # 2. Campos Opcionales con validación
    cedula_raw = str(data.get("cedula", "") or "").strip()
    # Clean dots/spaces from cedula
    cedula_clean = re.sub(r"[^\dVEve-]", "", cedula_raw)
    cedula_valid = cedula_clean if _REGEX_CEDULA.match(cedula_clean) else None

    edad_raw = data.get("edad")
    edad_valid = None
    if edad_raw:
        try:
            # Extraer solo números de la cadena por si dice "45 años"
            num_str = re.sub(r"[^\d]", "", str(edad_raw))
            if num_str:
                val = int(num_str)
                if 1 <= val <= 120:
                    edad_valid = val
        except (ValueError, TypeError):
            pass

    contacto_raw = str(data.get("contacto", "") or "").strip()
    contacto_valid = contacto_raw if _REGEX_TELEFONO.match(contacto_raw) else None

    sector = str(data.get("sector", "") or "").strip() or None

    return {
        "nombre_apellido": nombre_completo,
        "cedula": cedula_valid,
        "edad": edad_valid,
        "sector_ciudad": sector or "No especificado",
        "centro_salud": hospital,
        "telefono_contacto": contacto_valid or "No especificado",
        "quien_registra": "Carga Masiva",
    }, None


# ---------------------------------------------------------------------------
# Procesador: Excel (.xlsx / .xls)
# ---------------------------------------------------------------------------
def process_excel(filepath: str) -> Tuple[List[Dict], List[str]]:
    """Lee un archivo Excel usando pandas y extrae registros."""
    records = []
    errors = []

    try:
        df = pd.read_excel(filepath, engine="openpyxl" if filepath.endswith(".xlsx") else None)
    except Exception as e:
        return [], [f"Error al leer el archivo Excel: {e}"]

    if df.empty:
        return [], ["El archivo Excel está vacío."]

    col_mapping = _map_columns(df.columns.tolist())

    if not col_mapping:
        return [], [
            "No se detectaron columnas válidas. "
            "Asegúrese de incluir 'Nombre', 'Apellido' y 'Hospital'."
        ]

    for idx, row in df.iterrows():
        row_num = idx + 2  # Excel es 1-indexed + encabezado
        data = {}
        for orig_col, field in col_mapping.items():
            val = row.get(orig_col)
            if pd.notna(val):
                data[field] = val
            else:
                data[field] = ""

        record, error = _validate_record(data, row_num)
        if record:
            records.append(record)
        elif error:
            errors.append(error)

    return records, errors


# ---------------------------------------------------------------------------
# Extracción de texto con patrones (para PDF/Word no estructurado)
# ---------------------------------------------------------------------------
# Expresiones regulares sugeridas (priorizando Nombre y Hospital)
_TEXT_PATTERNS = {
    "nombre":   re.compile(r"(?i)nombre[s]?\s*[:=\-]\s*(?P<val>[a-záéíóúñA-ZÁÉÍÓÚÑ\s]+)"),
    "apellido": re.compile(r"(?i)apellido[s]?\s*[:=\-]\s*(?P<val>[a-záéíóúñA-ZÁÉÍÓÚÑ\s]+)"),
    "hospital": re.compile(r"(?i)(?:hospital|centro(?:\s+de\s+salud)?|cl[ií]nica)\s*[:=\-]\s*(?P<val>[^\n]+)"),
    "cedula":   re.compile(r"(?i)(?:c[eé]dula|c\.?i\.?|identidad)\s*[:=\-]\s*(?P<val>(?:[VEve][- ]?)?\d{6,9})"),
    "edad":     re.compile(r"(?i)(?:edad|age)\s*[:=\-]\s*(?P<val>\d{1,3})"),
    "contacto": re.compile(r"(?i)(?:contacto|tel[eé]fono|celular)\s*[:=\-]\s*(?P<val>[\d\+\-\s\(\)]{10,15})"),
    "sector":   re.compile(r"(?i)(?:sector|ciudad|ubicaci[oó]n)\s*[:=\-]\s*(?P<val>[^\n]+)"),
}

_SEPARATOR = re.compile(r"\n\s*\n|\n-{3,}\n|\n={3,}\n|\n\*{3,}\n")


def _parse_text_patterns(text: str) -> List[Dict]:
    """Divide el texto en bloques y busca los patrones solicitados."""
    results = []
    blocks = _SEPARATOR.split(text)

    for block in blocks:
        block = block.strip()
        if not block:
            continue

        data = {}
        for field, pattern in _TEXT_PATTERNS.items():
            match = pattern.search(block)
            if match:
                data[field] = match.group("val").strip()

        # Si encontramos al menos un indicio, es un candidato
        if data.get("nombre") or data.get("hospital"):
            results.append(data)

    return results


# ---------------------------------------------------------------------------
# Extracción con Gemini (Word / PDF)
# ---------------------------------------------------------------------------
def _extract_with_gemini(text: str) -> List[Dict] | None:
    """Intenta extraer registros de pacientes usando la API REST de Gemini."""
    import os
    import json
    import requests as http_req

    api_key = os.environ.get("GEMINI_API_KEY")
    if not api_key:
        return None

    try:
        prompt = f"""
Analiza el siguiente texto extraído de un documento de registro de pacientes afectados por la emergencia.
Tu tarea es extraer a todas las personas y devolver la información en formato JSON siguiendo estrictamente el esquema especificado.

Texto extraído:
---
{text}
---

Instrucciones:
1. Extrae todos los pacientes de las listas o tablas.
2. Identifica los siguientes campos para cada persona:
   - "nombre_apellido": Nombre completo (Nombres y Apellidos).
   - "hospital": Nombre del centro de salud u hospital donde se encuentra. Si no se especifica en la fila, búscalo en las cabeceras, títulos o pies de página de la misma sección (ej. Hospital Domingo Luciani, Cruz Roja, etc.).
   - "edad": Edad (entero, o nulo si no se especifica).
   - "cedula": Cédula de identidad o ID (limpia, solo números, sin puntos ni espacios; nulo si no se especifica).
   - "contacto": Número de teléfono de contacto (cadena de texto; nulo si no se especifica).
   - "sector": Dirección, sector o ciudad de procedencia (cadena de texto; "No especificado" si no se indica).
3. Resuelve cualquier split de columnas o páginas: si la información de una misma persona (por su nombre y número) está distribuida en distintas páginas (ej. nombre en página 1, cédula en página 9), únelas en un único registro completo.
4. Devuelve ÚNICAMENTE un objeto JSON válido con una clave "pacientes" que sea una lista de objetos con los campos descritos. NO AÑADAS TEXTO NI MARKDOWN. SOLO EL JSON PURO.

Esquema JSON esperado:
{{
  "pacientes": [
    {{
      "nombre_apellido": "Juan Perez",
      "hospital": "Hospital Domingo Luciani",
      "edad": 45,
      "cedula": "12345678",
      "contacto": "04121234567",
      "sector": "Catia"
    }}
  ]
}}
"""

        api_url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.0-flash:generateContent?key={{api_key}}"
        
        payload = {{
            "contents": [{{
                "parts": [{{"text": prompt}}]
            }}],
            "generationConfig": {{
                "responseMimeType": "application/json"
            }}
        }}

        resp = http_req.post(api_url, json=payload, timeout=90)
        
        if resp.status_code != 200:
            print(f"[Gemini Error] API HTTP {{resp.status_code}}: {{resp.text[:200]}}")
            return None

        result = resp.json()
        raw = result["candidates"][0]["content"]["parts"][0]["text"].strip()
        
        # Limpiar posibles bloques markdown
        if raw.startswith("```"):
            raw = raw.split("\n", 1)[-1].rsplit("```", 1)[0].strip()

        res_data = json.loads(raw)
        pacientes = res_data.get("pacientes", [])
        
        records = []
        for p in pacientes:
            data = {{
                "nombre_completo": str(p.get("nombre_apellido") or ""),
                "edad": p.get("edad"),
                "cedula": str(p.get("cedula") or ""),
                "contacto": str(p.get("contacto") or ""),
                "hospital": str(p.get("hospital") or ""),
                "sector": str(p.get("sector") or "")
            }}
            record, error = _validate_record(data, "Gemini API")
            if record:
                records.append(record)
        return records if records else None
    except Exception as e:
        print(f"[Gemini Error] Falló la extracción con Gemini: {{e}}")
        return None




# ---------------------------------------------------------------------------
# Procesador: Word (.docx)
# ---------------------------------------------------------------------------
def process_word(filepath: str) -> Tuple[List[Dict], List[str]]:
    """Extrae datos desde Tablas de Word o usando Patrones de texto."""
    records = []
    errors = []

    try:
        doc = DocxDocument(filepath)
    except Exception as e:
        return [], [f"Error al leer el archivo Word: {e}"]

    # Extraer texto de Word
    try:
        full_text = "\n".join(p.text for p in doc.paragraphs)
        for table in doc.tables:
            for row in table.rows:
                full_text += "\n" + " | ".join(cell.text.strip() for cell in row.cells)
                
        # 1. Intentar con la API de Gemini si la clave está configurada
        gemini_records = _extract_with_gemini(full_text)
        if gemini_records is not None:
            return gemini_records, []
    except Exception as e:
        print(f"Error al parsear Word o usar Gemini: {e}")

    # Lógica de extracción (Tablas y Texto)
    for table in doc.tables:
        if len(table.rows) < 2:
            continue
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        col_mapping = _map_columns(headers)
        if not col_mapping:
            continue

        for row_idx, row in enumerate(table.rows[1:], start=2):
            data = {}
            for col_idx, cell in enumerate(row.cells):
                if col_idx < len(headers):
                    orig_col = headers[col_idx]
                    if orig_col in col_mapping:
                        data[col_mapping[orig_col]] = cell.text.strip()

            record, error = _validate_record(data, row_idx)
            if record:
                records.append(record)
            elif error:
                errors.append(error)

    # Fallback local - Estrategia 2: Texto plano
    if not records and not doc.tables:
        full_text = "\n".join(p.text for p in doc.paragraphs)
        parsed = _parse_text_patterns(full_text)
        for idx, data in enumerate(parsed, start=1):
            record, error = _validate_record(data, idx)
            if record:
                records.append(record)
            elif error:
                errors.append(error)

    if not records and not errors:
        errors.append("El archivo Word no contiene tablas ni el formato 'Campo: Valor'.")

    return records, errors


# ---------------------------------------------------------------------------
# Procesador: PDF (.pdf)
# ---------------------------------------------------------------------------
def process_pdf(filepath: str) -> Tuple[List[Dict], List[str]]:
    """Extrae texto de un PDF y busca patrones de datos (RegEx) o tablas."""
    records = []
    errors = []

    try:
        reader = PdfReader(filepath)
    except Exception as e:
        return [], [f"Error al leer el archivo PDF: {e}"]

    if len(reader.pages) == 0:
        return [], ["El archivo PDF está vacío."]

    full_text = ""
    for page in reader.pages:
        text = page.extract_text()
        if text:
            full_text += text + "\n"

    # 1. Intentar con la API de Gemini si la clave está configurada
    try:
        gemini_records = _extract_with_gemini(full_text)
        if gemini_records is not None:
            return gemini_records, []
    except Exception as ge:
        print(f"Fallback local para PDF por error en Gemini: {ge}")

    # Extraer texto y procesar con expresiones regulares / parser local
    row_pat_with_hospital = re.compile(
        r"^\s*(\d+)\s+(Hospital\s+Universitario\s+de\s+Caracas|Hospital\s+Domingo\s+Luciani|Hospital\s+P[eé]rez\s+Carre[nñ]o|Cruz\s+Roja|Perif[eé]rico\s+de\s+Catia)\s+([a-zA-ZÁÉÍÓÚÑáéíóúñ\s\(\)/'\-]+?)\s+(\d+)(?:\s+([\d\.VE\-]+))?(?:\s+([\d\-]+))?",
        re.IGNORECASE
    )
    row_pat_without_hospital = re.compile(
        r"^\s*(\d+)\s+([a-zA-ZÁÉÍÓÚÑáéíóúñ\s\(\)/'\-]+?)\s+(\d+)(?:\s+([\d\.VE\-]+))?(?:\s+([\d\-]+))?",
        re.IGNORECASE
    )

    current_hospital = None
    
    # Reset full_text scan page by page
    for page_idx, page in enumerate(reader.pages, start=1):
        text = page.extract_text()
        if not text:
            continue
        lines = text.split("\n")
        
        page_hospital = None
        for line in lines:
            line_clean = line.strip()
            if re.match(r"^(?:HOSPITAL|CDI|AMBULATORIO|CL[IÍ]NICA|CRUZ ROJA|PERIF[EÉ]RICO)\b", line_clean, re.IGNORECASE) and not re.match(r"^\d", line_clean):
                page_hospital = line_clean
                current_hospital = line_clean
                break

        for line in lines:
            line_clean = line.strip()
            
            m = row_pat_with_hospital.match(line_clean)
            if m:
                hosp = m.group(2).strip()
                name = m.group(3).strip()
                age = m.group(4).strip()
                cedula = m.group(5).strip() if m.group(5) else ""
                telefono = m.group(6).strip() if m.group(6) else ""
                
                data = {
                    "nombre_completo": name,
                    "edad": age,
                    "cedula": cedula,
                    "contacto": telefono,
                    "hospital": hosp
                }
                record, error = _validate_record(data, f"Pág {page_idx} - Fila {m.group(1)}")
                if record:
                    records.append(record)
                elif error:
                    errors.append(error)
                continue

            m = row_pat_without_hospital.match(line_clean)
            if m:
                name = m.group(2).strip()
                age = m.group(3).strip()
                cedula = m.group(4).strip() if m.group(4) else ""
                telefono = m.group(5).strip() if m.group(5) else ""
                hosp = page_hospital or current_hospital or "Desconocido"

                data = {
                    "nombre_completo": name,
                    "edad": age,
                    "cedula": cedula,
                    "contacto": telefono,
                    "hospital": hosp
                }
                record, error = _validate_record(data, f"Pág {page_idx} - Fila {m.group(1)}")
                if record:
                    records.append(record)
                elif error:
                    errors.append(error)

    if not records:
        parsed = _parse_text_patterns(full_text)
        for idx, data in enumerate(parsed, start=1):
            record, error = _validate_record(data, idx)
            if record:
                records.append(record)
            elif error:
                errors.append(error)

    if not records and not errors:
        errors.append("No se pudo detectar el formato de datos en el PDF. Asegúrese de incluir una tabla o el formato 'Campo: Valor'.")

    return records, errors





# ---------------------------------------------------------------------------
# Dispatcher
# ---------------------------------------------------------------------------
def procesar_archivo(file_path: str, extension: str) -> Tuple[List[Dict], List[str]]:
    """
    Punto de entrada principal.
    Redirige al procesador correspondiente y aplica las reglas de extracción.
    Retorna (records, errors).
    """
    extension = extension.lower().strip(".")

    if extension in ("xlsx", "xls"):
        return process_excel(file_path)
    elif extension == "docx":
        return process_word(file_path)
    elif extension == "pdf":
        return process_pdf(file_path)
    else:
        return [], [f"Formato no soportado: .{extension}"]
