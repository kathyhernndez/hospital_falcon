"""
Aplicacion Flask -- Registro de Personas en Centros de Salud de Venezuela
=========================================================================
Permite registrar y buscar personas ubicadas en hospitales y centros de
emergencia a lo largo del territorio venezolano.

Incluye:
- Sistema de autenticacion para personal de salud (Flask-Login)
- Carga masiva de datos desde archivos Excel, Word y PDF
- Generacion automatica de plantillas de descarga
"""

import os
import re
from datetime import datetime, timezone, timedelta

# Load variables from .env file if present
if os.path.exists(".env"):
    with open(".env", "r", encoding="utf-8") as f:
        for line in f:
            line_strip = line.strip()
            if line_strip and not line_strip.startswith("#") and "=" in line_strip:
                key, val = line_strip.split("=", 1)
                os.environ[key.strip()] = val.strip().strip('"').strip("'")

from flask import (
    Flask, render_template, request, jsonify, redirect, url_for, flash,
    send_file, session,
)
from flask_login import (
    LoginManager, login_user, logout_user, login_required, current_user,
)
from werkzeug.utils import secure_filename

from models import db, Registro, User, HealthCenter
from file_processors import procesar_archivo

# ---------------------------------------------------------------------------
# Configuracion
# ---------------------------------------------------------------------------
BASE_DIR = os.path.abspath(os.path.dirname(__file__))
UPLOAD_FOLDER = os.path.join(BASE_DIR, "static", "uploads")
BULK_UPLOAD_FOLDER = os.path.join(BASE_DIR, "static", "bulk_uploads")
TEMPLATE_FOLDER_DL = os.path.join(BASE_DIR, "static", "templates")
ID_UPLOAD_FOLDER = os.path.join(BASE_DIR, "private", "ids")
ALLOWED_EXTENSIONS = {"png", "jpg", "jpeg", "gif", "webp"}
ALLOWED_BULK_EXTENSIONS = {"xlsx", "xls", "docx", "pdf"}
ALLOWED_ID_EXTENSIONS = {"png", "jpg", "jpeg", "pdf"}

app = Flask(__name__)
app.config["SQLALCHEMY_DATABASE_URI"] = f"sqlite:///{os.path.join(BASE_DIR, 'hospitales.db')}"
app.config["SQLALCHEMY_TRACK_MODIFICATIONS"] = False
app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER
app.config["ID_UPLOAD_FOLDER"] = ID_UPLOAD_FOLDER
app.config["MAX_CONTENT_LENGTH"] = 16 * 1024 * 1024  # 16 MB maximo
app.config["SECRET_KEY"] = "hospitales-ve-secret-key-2026-cambiar-en-produccion"

db.init_app(app)

# --- Flask-Login ---
login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = "login"
login_manager.login_message = "Debe iniciar sesion para acceder a esta seccion."
login_manager.login_message_category = "warning"


@login_manager.user_loader
def load_user(user_id):
    return db.session.get(User, int(user_id))


# Crear directorios necesarios
for folder in [UPLOAD_FOLDER, BULK_UPLOAD_FOLDER, TEMPLATE_FOLDER_DL, ID_UPLOAD_FOLDER]:
    os.makedirs(folder, exist_ok=True)


# ---------------------------------------------------------------------------
# Utilidades
# ---------------------------------------------------------------------------
def allowed_file(filename: str) -> bool:
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS


def allowed_bulk_file(filename: str) -> bool:
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_BULK_EXTENSIONS


def allowed_id_file(filename: str) -> bool:
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_ID_EXTENSIONS


def generate_templates():
    """Genera las plantillas de descarga (Excel y Word) si no existen."""
    excel_path = os.path.join(TEMPLATE_FOLDER_DL, "plantilla_pacientes.xlsx")
    word_path = os.path.join(TEMPLATE_FOLDER_DL, "plantilla_pacientes.docx")

    # --- Plantilla Excel ---
    if not os.path.exists(excel_path):
        try:
            from openpyxl import Workbook
            from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

            wb = Workbook()
            ws = wb.active
            ws.title = "Pacientes"

            headers = ["Nombre", "Apellido", "Cedula", "Edad", "Sector", "Hospital", "Contacto"]
            header_fill = PatternFill(start_color="059669", end_color="059669", fill_type="solid")
            header_font = Font(name="Arial", bold=True, color="FFFFFF", size=11)
            thin_border = Border(
                left=Side(style="thin"), right=Side(style="thin"),
                top=Side(style="thin"), bottom=Side(style="thin"),
            )

            for col_idx, header in enumerate(headers, 1):
                cell = ws.cell(row=1, column=col_idx, value=header)
                cell.font = header_font
                cell.fill = header_fill
                cell.alignment = Alignment(horizontal="center")
                cell.border = thin_border

            # Fila de ejemplo
            example = ["Maria", "Garcia Lopez", "V-12345678", 45, "Caracas, San Juan", "Hospital Universitario de Caracas", "0412-1234567"]
            for col_idx, val in enumerate(example, 1):
                cell = ws.cell(row=2, column=col_idx, value=val)
                cell.border = thin_border
                cell.alignment = Alignment(horizontal="center")

            # Ajustar anchos
            widths = [18, 20, 16, 8, 25, 40, 18]
            for i, w in enumerate(widths, 1):
                ws.column_dimensions[chr(64 + i)].width = w

            wb.save(excel_path)
        except Exception:
            pass

    # --- Plantilla Word ---
    if not os.path.exists(word_path):
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.table import WD_TABLE_ALIGNMENT

            doc = Document()

            # Titulo
            title = doc.add_heading("Plantilla de Registro de Pacientes", level=1)
            for run in title.runs:
                run.font.color.rgb = RGBColor(0x05, 0x96, 0x69)

            doc.add_paragraph(
                "Complete la siguiente tabla con los datos de los pacientes. "
                "Las columnas Nombre y Hospital son obligatorias."
            )

            headers = ["Nombre", "Apellido", "Cedula", "Edad", "Sector", "Hospital", "Contacto"]
            table = doc.add_table(rows=2, cols=len(headers))
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            # Encabezados
            for idx, header in enumerate(headers):
                cell = table.rows[0].cells[idx]
                cell.text = header
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(10)

            # Fila de ejemplo
            example = ["Maria", "Garcia Lopez", "V-12345678", "45", "Caracas", "Hospital Universitario de Caracas", "0412-1234567"]
            for idx, val in enumerate(example):
                table.rows[1].cells[idx].text = val

            doc.save(word_path)
        except Exception:
            pass


# ---------------------------------------------------------------------------
# Datos de Centros de Salud de Venezuela (precargados)
# ---------------------------------------------------------------------------
CENTROS_SALUD = [
    # -- Distrito Capital --
    {"nombre": "Hospital Universitario de Caracas", "lat": 10.4918, "lng": -66.8459, "estado": "Distrito Capital"},
    {"nombre": "Hospital Vargas de Caracas", "lat": 10.5065, "lng": -66.9125, "estado": "Distrito Capital"},
    {"nombre": "Hospital de Clinicas Caracas", "lat": 10.4965, "lng": -66.8560, "estado": "Distrito Capital"},
    {"nombre": "Hospital Militar Dr. Carlos Arvelo", "lat": 10.4950, "lng": -66.8780, "estado": "Distrito Capital"},
    {"nombre": "Hospital de Ninos J.M. de los Rios", "lat": 10.5032, "lng": -66.8914, "estado": "Distrito Capital"},
    {"nombre": "Maternidad Concepcion Palacios", "lat": 10.5010, "lng": -66.9190, "estado": "Distrito Capital"},
    {"nombre": "Hospital Domingo Luciani (El Llanito)", "lat": 10.4750, "lng": -66.8050, "estado": "Distrito Capital"},
    # -- Miranda --
    {"nombre": "Hospital Perez de Leon", "lat": 10.4780, "lng": -66.8160, "estado": "Miranda"},
    {"nombre": "Hospital General de Los Valles del Tuy", "lat": 10.2350, "lng": -66.9800, "estado": "Miranda"},
    # -- Zulia --
    {"nombre": "Hospital Universitario de Maracaibo", "lat": 10.6427, "lng": -71.6370, "estado": "Zulia"},
    {"nombre": "Hospital Central Dr. Urquinaona", "lat": 10.6490, "lng": -71.6340, "estado": "Zulia"},
    {"nombre": "Hospital General del Sur Dr. Pedro Iturbe", "lat": 10.6200, "lng": -71.6500, "estado": "Zulia"},
    {"nombre": "Hospital de Especialidades Pediatricas", "lat": 10.6380, "lng": -71.6290, "estado": "Zulia"},
    # -- Carabobo --
    {"nombre": "Ciudad Hospitalaria Dr. Enrique Tejera", "lat": 10.1810, "lng": -68.0050, "estado": "Carabobo"},
    {"nombre": "Hospital Central de Valencia", "lat": 10.1795, "lng": -67.9960, "estado": "Carabobo"},
    {"nombre": "Hospital Metropolitano del Norte", "lat": 10.2100, "lng": -67.9800, "estado": "Carabobo"},
    # -- Lara --
    {"nombre": "Hospital Central Antonio Maria Pineda", "lat": 10.0678, "lng": -69.3220, "estado": "Lara"},
    {"nombre": "Hospital Pediatrico Dr. Agustin Zubillaga", "lat": 10.0650, "lng": -69.3180, "estado": "Lara"},
    # -- Aragua --
    {"nombre": "Hospital Central de Maracay", "lat": 10.2470, "lng": -67.5960, "estado": "Aragua"},
    {"nombre": "Hospital Jose Maria Benitez", "lat": 10.1700, "lng": -67.4600, "estado": "Aragua"},
    # -- Merida --
    {"nombre": "Hospital Universitario de Los Andes (IAHULA)", "lat": 8.5897, "lng": -71.1561, "estado": "Merida"},
    # -- Bolivar --
    {"nombre": "Hospital Ruiz y Paez", "lat": 8.3550, "lng": -62.6410, "estado": "Bolivar"},
    {"nombre": "Hospital Uyapar", "lat": 8.2880, "lng": -62.7350, "estado": "Bolivar"},
    # -- Tachira --
    {"nombre": "Hospital Central de San Cristobal", "lat": 7.7669, "lng": -72.2250, "estado": "Tachira"},
    {"nombre": "Hospital Dr. Patrocinio Penuela Ruiz", "lat": 7.7700, "lng": -72.2200, "estado": "Tachira"},
    # -- Anzoategui --
    {"nombre": "Hospital Universitario Dr. Luis Razetti", "lat": 10.1370, "lng": -64.6880, "estado": "Anzoategui"},
    {"nombre": "Hospital Central de Barcelona", "lat": 10.1300, "lng": -64.7000, "estado": "Anzoategui"},
    # -- Falcon --
    {"nombre": "Hospital Universitario Dr. Alfredo Van Grieken", "lat": 11.4160, "lng": -69.6730, "estado": "Falcon"},
    # -- Sucre --
    {"nombre": "Hospital Universitario Antonio Patricio de Alcala", "lat": 10.4600, "lng": -64.1770, "estado": "Sucre"},
    # -- Portuguesa --
    {"nombre": "Hospital Miguel Oraa", "lat": 9.0590, "lng": -69.7430, "estado": "Portuguesa"},
    # -- Barinas --
    {"nombre": "Hospital Luis Razetti de Barinas", "lat": 8.6280, "lng": -70.2070, "estado": "Barinas"},
    # -- Monagas --
    {"nombre": "Hospital Universitario Dr. Manuel Nunez Tovar", "lat": 9.7450, "lng": -63.1830, "estado": "Monagas"},
    # -- Trujillo --
    {"nombre": "Hospital Central de Valera", "lat": 9.3170, "lng": -70.6030, "estado": "Trujillo"},
    # -- Yaracuy --
    {"nombre": "Hospital Central de San Felipe", "lat": 10.3400, "lng": -68.7400, "estado": "Yaracuy"},
    # -- Nueva Esparta --
    {"nombre": "Hospital Central Dr. Luis Ortega", "lat": 11.0000, "lng": -63.8600, "estado": "Nueva Esparta"},
    # -- Guarico --
    {"nombre": "Hospital Israel Ranuarez Balza", "lat": 9.9060, "lng": -67.3560, "estado": "Guarico"},
    # -- Apure --
    {"nombre": "Hospital Pablo Acosta Ortiz", "lat": 7.8830, "lng": -67.4720, "estado": "Apure"},
    # -- Cojedes --
    {"nombre": "Hospital General de San Carlos", "lat": 9.6590, "lng": -68.5880, "estado": "Cojedes"},
    # -- Delta Amacuro --
    {"nombre": "Hospital Razetti de Tucupita", "lat": 9.0580, "lng": -62.0500, "estado": "Delta Amacuro"},
    # -- Amazonas --
    {"nombre": "Hospital Dr. Jose Gregorio Hernandez", "lat": 5.6630, "lng": -67.6270, "estado": "Amazonas"},
    # -- Vargas (La Guaira) --
    {"nombre": "Hospital Jose Maria Vargas de La Guaira", "lat": 10.6020, "lng": -66.9340, "estado": "Vargas"},
]


# ---------------------------------------------------------------------------
# Rutas: Pagina Principal y API de Busqueda
# ---------------------------------------------------------------------------
@app.route("/")
def index():
    """Pagina principal -- Dashboard de busqueda."""
    registros = Registro.query.order_by(Registro.fecha_registro.desc()).all()
    total = Registro.query.count()
    total_centers = HealthCenter.query.filter_by(is_approved=True).count()
    
    # Serialize registers for initial JS state
    import json
    registros_json = json.dumps([r.to_dict(mask_phone=False) for r in registros])
    
    return render_template(
        "index.html",
        registros=registros,
        total=total,
        total_centers=total_centers,
        registros_json=registros_json
    )


@app.route("/api/search")
def api_search():
    """Busca registros por nombre, cedula o centro de salud."""
    q = request.args.get("q", "").strip()
    filtro = request.args.get("filter", "nombre")

    query = Registro.query

    if q:
        if filtro == "cedula":
            query = query.filter(Registro.cedula.ilike(f"%{q}%"))
        elif filtro == "centro":
            query = query.join(HealthCenter).filter(HealthCenter.nombre.ilike(f"%{q}%"))
        else:  # nombre (por defecto)
            query = query.filter(Registro.nombre_apellido.ilike(f"%{q}%"))

    registros = query.order_by(Registro.fecha_registro.desc()).all()
    results = [r.to_dict(mask_phone=False) for r in registros]

    return jsonify({"registros": results, "total": len(results)})


@app.route("/api/registrar", methods=["POST"])
def api_registrar():
    """Registra una nueva persona en un centro de salud."""
    try:
        # Campos obligatorios
        nombre = request.form.get("nombre_apellido", "").strip()
        sector = request.form.get("sector_ciudad", "").strip()
        quien = request.form.get("quien_registra", "").strip()
        telefono = request.form.get("telefono_contacto", "").strip()

        # Validacion centro de salud
        health_center_id_str = request.form.get("health_center_id", "").strip()
        centro_nuevo_nombre = request.form.get("nuevo_centro_nombre", "").strip()
        centro_nuevo_tipo = request.form.get("nuevo_centro_tipo", "").strip()

        if not all([nombre, sector, quien, telefono]):
            return jsonify({"error": "Todos los campos obligatorios deben ser completados."}), 400

        # Sanitizar y validar Nombre y Apellido: min 2 max 20, letras únicamente
        if not re.match(r"^[a-zA-ZáéíóúÁÉÍÓÚñÑ\s]{2,20}$", nombre):
            return jsonify({"error": "El nombre y apellido debe contener únicamente letras y tener entre 2 y 20 caracteres."}), 400

        # Sanitizar y validar teléfono: min 12 números
        telefono_clean = re.sub(r"\D", "", telefono)
        if len(telefono_clean) < 12:
            return jsonify({"error": "El teléfono de contacto debe contener al menos 12 dígitos numéricos."}), 400
        telefono = telefono_clean

        health_center_id = None
        if health_center_id_str:
            health_center_id = int(health_center_id_str)
        elif centro_nuevo_nombre and centro_nuevo_tipo:
            if not re.match(r"^[a-zA-ZáéíóúÁÉÍÓÚñÑ\s]{8,50}$", centro_nuevo_nombre):
                return jsonify({"error": "El nombre del nuevo centro de salud debe contener únicamente letras y tener entre 8 y 50 caracteres."}), 400
            # Crear nuevo centro no aprobado
            nuevo_centro = HealthCenter(
                nombre=centro_nuevo_nombre,
                tipo=centro_nuevo_tipo,
                is_approved=False
            )
            db.session.add(nuevo_centro)
            db.session.commit() # para obtener el ID
            health_center_id = nuevo_centro.id
        else:
             return jsonify({"error": "Debe seleccionar un centro de salud o proporcionar uno nuevo."}), 400

        # Campos opcionales
        cedula = request.form.get("cedula", "").strip() or None
        if cedula:
            cedula_clean = re.sub(r"\D", "", cedula)
            if not re.match(r"^\d{7,8}$", cedula_clean):
                return jsonify({"error": "La cédula debe contener únicamente números y tener entre 7 y 8 dígitos."}), 400
            cedula = f"V-{cedula_clean}"

        edad_str = request.form.get("edad", "").strip()
        edad = None
        if edad_str:
            try:
                edad = int(edad_str)
                if edad < 1 or edad > 99:
                    raise ValueError()
            except ValueError:
                return jsonify({"error": "La edad debe ser un número entero entre 1 y 99."}), 400
        trabaja = request.form.get("trabaja_centro") == "true"

        # Manejo de foto
        foto_path = None
        if "foto" in request.files:
            archivo = request.files["foto"]
            if archivo and archivo.filename and allowed_file(archivo.filename):
                nombre_seguro = secure_filename(archivo.filename)
                base, ext = os.path.splitext(nombre_seguro)
                nombre_unico = f"{base}_{int(datetime.now(timezone.utc).timestamp())}{ext}"
                archivo.save(os.path.join(app.config["UPLOAD_FOLDER"], nombre_unico))
                foto_path = nombre_unico

        force_transfer = request.form.get("force_transfer") == "true"

        if cedula:
            existing_record = Registro.query.filter_by(cedula=cedula).first()
            if existing_record:
                if existing_record.health_center_id == health_center_id:
                    return jsonify({"error": "Usuario ya registrado en este centro médico."}), 400
                else:
                    old_center_name = existing_record.health_center.nombre if existing_record.health_center else "Centro Desconocido"
                    new_center = db.session.get(HealthCenter, health_center_id)
                    new_center_name = new_center.nombre if new_center else "Centro Nuevo"
                    
                    if force_transfer:
                        existing_record.health_center_id = health_center_id
                        
                        transfer_note = f"Origen: {old_center_name}\nTrasladado a: {new_center_name}"
                        if existing_record.observaciones:
                            existing_record.observaciones += f"\n\n{transfer_note}"
                        else:
                            existing_record.observaciones = transfer_note
                            
                        db.session.commit()
                        return jsonify({"success": True, "message": "Paciente trasladado exitosamente.", "id": existing_record.id})
                    else:
                        return jsonify({
                            "requires_transfer": True, 
                            "existing_center": old_center_name,
                            "new_center": new_center_name
                        })

        # Crear registro si no existe duplicado
        registro = Registro(
            nombre_apellido=nombre,
            sector_ciudad=sector,
            health_center_id=health_center_id,
            quien_registra=quien,
            telefono_contacto=telefono,
            cedula=cedula,
            edad=edad,
            trabaja_centro=trabaja,
            foto_path=foto_path,
        )

        db.session.add(registro)
        db.session.commit()

        return jsonify({"success": True, "message": "Registro exitoso.", "id": registro.id})

    except ValueError as ve:
        db.session.rollback()
        return jsonify({"error": f"Dato invalido: {ve}"}), 400
    except Exception as e:
        db.session.rollback()
        return jsonify({"error": f"Error interno: {e}"}), 500


@app.route("/api/centros-salud")
def api_centros_salud():
    """Devuelve la lista de centros de salud aprobados."""
    centros = HealthCenter.query.filter_by(is_approved=True).all()
    return jsonify([c.to_dict() for c in centros])


# ---------------------------------------------------------------------------
# Rutas: Autenticacion (Login / Register / Logout)
# ---------------------------------------------------------------------------
@app.route("/login", methods=["GET", "POST"])
def login():
    """Inicio de sesion para personal de salud."""
    if current_user.is_authenticated:
        return redirect(url_for("upload_dashboard"))

    if request.method == "POST":
        correo = request.form.get("correo", "").strip().lower()
        password = request.form.get("password", "").strip()

        if not correo or not password:
            flash("Por favor complete todos los campos.", "error")
            return render_template("login.html")

        # Sanitizar y validar correo: formato básico con @ y dominio
        if not re.match(r'^[^@\s]+@[^@\s]+\.[^@\s]+$', correo):
            flash("El correo electrónico no tiene un formato válido.", "error")
            return render_template("login.html")

        user = User.query.filter_by(correo=correo).first()

        if user:
            # Check if account is locked out
            if user.lockout_until:
                # Make timezone-aware if needed
                lockout_until = user.lockout_until
                if lockout_until.tzinfo is None:
                    lockout_until = lockout_until.replace(tzinfo=timezone.utc)
                
                now = datetime.now(timezone.utc)
                if lockout_until > now:
                    remaining = int((lockout_until - now).total_seconds())
                    minutes = remaining // 60
                    seconds = remaining % 60
                    time_str = f"{minutes} min {seconds} s" if minutes > 0 else f"{seconds} s"
                    flash(f"Su cuenta está temporalmente bloqueada por exceso de intentos fallidos. Intente de nuevo en {time_str}.", "error")
                    return render_template("login.html")

            if user.check_password(password):
                # Reset attempts on success
                user.failed_login_attempts = 0
                user.lockout_until = None
                db.session.commit()

                if user.account_status == "pending":
                    flash("Su cuenta está en proceso de verificación por un administrador. No puede iniciar sesión aún.", "warning")
                    return render_template("login.html")
                elif user.account_status == "restricted":
                    flash("Su cuenta ha sido restringida. Contacte al administrador.", "error")
                    return render_template("login.html")
                    
                login_user(user, remember=True)
                flash(f"Bienvenido/a, {user.nombre}!", "success")
                
                if user.role == "admin":
                    return redirect(url_for("admin_dashboard"))
                    
                next_page = request.args.get("next")
                return redirect(next_page or url_for("upload_dashboard"))
            else:
                user.failed_login_attempts += 1
                if user.failed_login_attempts >= 3:
                    user.lockout_until = datetime.now(timezone.utc) + timedelta(minutes=5)
                    user.failed_login_attempts = 0
                    db.session.commit()
                    flash("Su cuenta ha sido bloqueada por 5 minutos debido a 3 intentos fallidos de inicio de sesión.", "error")
                else:
                    db.session.commit()
                    flash(f"Correo o contraseña incorrectos. Le quedan {3 - user.failed_login_attempts} intentos antes de ser bloqueado por 5 minutos.", "error")
        else:
            flash("Correo o contraseña incorrectos.", "error")

    return render_template("login.html")


@app.route("/register", methods=["GET", "POST"])
def register():
    """Registro de cuenta para personal de salud."""
    if current_user.is_authenticated:
        return redirect(url_for("upload_dashboard"))

    centros = HealthCenter.query.filter_by(is_approved=True).order_by(HealthCenter.nombre.asc()).all()

    if request.method == "POST":
        nombre = request.form.get("nombre", "").strip()
        correo = request.form.get("correo", "").strip().lower()
        password = request.form.get("password", "").strip()
        password_confirm = request.form.get("password_confirm", "").strip()
        centro = request.form.get("centro_medico", "").strip()

        # Check if they suggested a new center
        toggle_nuevo = request.form.get("toggle_nuevo_centro") == "on"
        nuevo_centro_nombre = request.form.get("nuevo_centro_nombre", "").strip()
        nuevo_centro_tipo = request.form.get("nuevo_centro_tipo", "").strip()
        nuevo_centro_lat = request.form.get("nuevo_centro_lat", "").strip()
        nuevo_centro_lng = request.form.get("nuevo_centro_lng", "").strip()

        # Validaciones
        errors = []
        if toggle_nuevo:
            if not all([nombre, correo, password, password_confirm, nuevo_centro_nombre, nuevo_centro_tipo, nuevo_centro_lat, nuevo_centro_lng]):
                errors.append("Todos los campos son obligatorios. Asegúrese de ubicar el nuevo centro de salud en el mapa.")
            # Validar nuevo centro: min 8 max 50, letras únicamente
            if not re.match(r"^[a-zA-ZáéíóúÁÉÍÓÚñÑ\s]{8,50}$", nuevo_centro_nombre):
                errors.append("El nombre del nuevo centro de salud debe contener únicamente letras y tener entre 8 y 50 caracteres.")
            centro = nuevo_centro_nombre
        else:
            if not all([nombre, correo, password, password_confirm, centro]):
                errors.append("Todos los campos son obligatorios.")

        # Sanitizar y validar Nombre y Apellido: min 2 max 20, letras únicamente
        if not re.match(r"^[a-zA-ZáéíóúÁÉÍÓÚñÑ\s]{2,20}$", nombre):
            errors.append("El nombre completo debe contener únicamente letras y tener entre 2 y 20 caracteres.")

        # Sanitizar y validar correo: formato básico con @ y dominio
        if not re.match(r'^[^@\s]+@[^@\s]+\.[^@\s]+$', correo):
            errors.append("El correo electrónico no tiene un formato válido.")

        if password != password_confirm:
            errors.append("Las contrasenas no coinciden.")
        if len(password) < 10 or len(password) > 16:
            errors.append("La contraseña debe tener entre 10 y 16 caracteres.")
        if not re.search(r"[A-Za-z]", password) or not re.search(r"\d", password):
            errors.append("La contraseña debe ser una combinación alfanumérica (letras y números).")
        if User.query.filter_by(correo=correo).first():
            errors.append("Ya existe una cuenta con ese correo electronico.")
            
        # Validar foto de carnet
        id_card_path = None
        if "id_card" not in request.files:
            errors.append("Debe subir una foto o PDF de su carnet institucional.")
        else:
            archivo = request.files["id_card"]
            if archivo.filename == "":
                errors.append("No seleccionó ningún archivo de carnet.")
            elif not allowed_id_file(archivo.filename):
                errors.append("Formato de carnet no permitido. Use .jpg, .png o .pdf.")
            else:
                import uuid
                ext = archivo.filename.rsplit(".", 1)[1].lower()
                nombre_seguro = f"{uuid.uuid4().hex}.{ext}"
                try:
                    archivo.save(os.path.join(app.config["ID_UPLOAD_FOLDER"], nombre_seguro))
                    id_card_path = nombre_seguro
                except Exception as e:
                    errors.append(f"Error al guardar el carnet: {e}")

        if errors:
            for err in errors:
                flash(err, "error")
            return render_template("register.html", centros=centros)

        # Crear nuevo centro no aprobado si corresponde
        if toggle_nuevo:
            try:
                lat = float(nuevo_centro_lat) if nuevo_centro_lat else None
                lng = float(nuevo_centro_lng) if nuevo_centro_lng else None
                nuevo_centro = HealthCenter(
                    nombre=nuevo_centro_nombre,
                    tipo=nuevo_centro_tipo,
                    latitud=lat,
                    longitud=lng,
                    is_approved=False
                )
                db.session.add(nuevo_centro)
                db.session.commit()
            except Exception as e:
                db.session.rollback()
                flash(f"Error al crear el nuevo centro médico: {e}", "error")
                return render_template("register.html", centros=centros)

        # Crear usuario
        user = User(
            nombre=nombre,
            correo=correo,
            centro_medico_asociado=centro,
            is_verified_staff=True,
            role="health_staff",
            account_status="pending",
            id_card_path=id_card_path
        )
        user.set_password(password)

        try:
            db.session.add(user)
            db.session.commit()
            flash("Registro completado. Su cuenta está en proceso de verificación por un administrador.", "success")
            return redirect(url_for("login"))
        except Exception:
            db.session.rollback()
            flash("Error al crear la cuenta. Intente de nuevo.", "error")

    return render_template("register.html", centros=centros)


@app.route("/logout")
@login_required
def logout():
    """Cierra la sesion del usuario."""
    logout_user()
    flash("Sesion cerrada correctamente.", "success")
    return redirect(url_for("index"))


def normalize_center_name(name):
    if not name:
        return ""
    n = name.lower().strip()
    accents = {'á': 'a', 'é': 'e', 'í': 'i', 'ó': 'o', 'ú': 'u', 'ü': 'u', 'ñ': 'n'}
    for char, replacement in accents.items():
        n = n.replace(char, replacement)
    for prefix in ["hospital universitario de ", "hospital universitario ", "hospital de ", "hospital ", "cdi de ", "cdi ", "ambulatorio de ", "ambulatorio ", "clinica ", "clínica "]:
        if n.startswith(prefix):
            n = n[len(prefix):]
            break
    n = re.sub(r"[^\w\s]", "", n)
    n = re.sub(r"\s+", " ", n).strip()
    return n


def find_existing_center(name):
    cleaned_name = re.sub(r"^\d+[-\s]+", "", name).strip()
    normalized_input = normalize_center_name(cleaned_name)
    if not normalized_input:
        return None
    all_centers = HealthCenter.query.all()
    for center in all_centers:
        if normalize_center_name(center.nombre) == normalized_input:
            return center
    return None


# ---------------------------------------------------------------------------
# Rutas: Panel de Carga Masiva (protegido)
# ---------------------------------------------------------------------------
@app.route("/dashboard/upload")
@login_required
def upload_dashboard():
    """Panel de carga masiva de datos (solo personal autenticado)."""
    if current_user.account_status != "approved" and current_user.role != "admin":
        flash("Su cuenta no está aprobada para realizar cargas masivas.", "warning")
        return redirect(url_for("index"))
    return render_template("upload.html")


@app.route("/api/upload-bulk", methods=["POST"])
@login_required
def api_upload_bulk():
    """Procesa un archivo subido (Excel, Word o PDF) y crea registros masivos."""
    if current_user.account_status != "approved" and current_user.role != "admin":
        return jsonify({"error": "Su cuenta no está aprobada para esta acción."}), 403
    if "archivo" not in request.files:
        return jsonify({"error": "No se selecciono ningun archivo."}), 400

    archivo = request.files["archivo"]

    if not archivo or not archivo.filename:
        return jsonify({"error": "Archivo vacio o sin nombre."}), 400

    if not allowed_bulk_file(archivo.filename):
        return jsonify({
            "error": "Formato no permitido. Use archivos .xlsx, .xls, .docx o .pdf."
        }), 400

    # Guardar archivo temporalmente
    nombre_seguro = secure_filename(archivo.filename)
    timestamp = int(datetime.now(timezone.utc).timestamp())
    nombre_unico = f"{timestamp}_{nombre_seguro}"
    filepath = os.path.join(BULK_UPLOAD_FOLDER, nombre_unico)
    archivo.save(filepath)

    try:
        # Procesar archivo
        ext = archivo.filename.rsplit(".", 1)[-1].lower() if "." in archivo.filename else ""
        records, errors = procesar_archivo(filepath, ext)

        # Insertar registros validos en la base de datos
        inserted = 0
        db_errors = []

        for rec in records:
            try:
                hospital_name = rec.get("centro_salud", "Desconocido").strip()
                # Clean leading numbers if any
                hospital_name_clean = re.sub(r"^\d+[-\s]+", "", hospital_name)
                
                # Check if center exists
                center = find_existing_center(hospital_name)
                if not center:
                    center = HealthCenter(
                        nombre=hospital_name_clean,
                        tipo="Hospital",
                        is_approved=False  # New health centers from bulk upload must be approved by admin
                    )
                    db.session.add(center)
                    db.session.commit()
                
                health_center_id = center.id

                # Deduplicate / Merge
                existing = Registro.query.filter(
                    db.func.lower(Registro.nombre_apellido) == rec["nombre_apellido"].lower(),
                    Registro.health_center_id == health_center_id
                ).first()

                if existing:
                    updated = False
                    if rec.get("cedula") and not existing.cedula:
                        existing.cedula = rec["cedula"]
                        updated = True
                    if rec.get("edad") and not existing.edad:
                        existing.edad = rec["edad"]
                        updated = True
                    if rec.get("telefono_contacto") and (not existing.telefono_contacto or existing.telefono_contacto == "No especificado"):
                        existing.telefono_contacto = rec["telefono_contacto"]
                        updated = True
                    if rec.get("sector_ciudad") and (not existing.sector_ciudad or existing.sector_ciudad == "No especificado"):
                        existing.sector_ciudad = rec["sector_ciudad"]
                        updated = True
                    
                    if updated:
                        db.session.commit()
                    continue

                registro = Registro(
                    nombre_apellido=rec["nombre_apellido"],
                    sector_ciudad=rec.get("sector_ciudad", "No especificado"),
                    health_center_id=health_center_id,
                    quien_registra=rec.get("quien_registra", "Carga Masiva"),
                    telefono_contacto=rec.get("telefono_contacto", "No especificado"),
                    cedula=rec.get("cedula"),
                    edad=rec.get("edad"),
                )
                db.session.add(registro)
                inserted += 1
            except Exception as e:
                db_errors.append(f"Error al insertar/actualizar '{rec.get('nombre_apellido', '?')}': {e}")

        # Commit en bloque
        try:
            db.session.commit()
        except Exception as e:
            db.session.rollback()
            return jsonify({
                "error": f"Error al guardar en la base de datos: {e}",
                "registros_procesados": 0,
                "errores": errors + db_errors,
            }), 500

        all_errors = errors + db_errors

        return jsonify({
            "success": True,
            "registros_procesados": inserted,
            "errores_count": len(all_errors),
            "errores": all_errors,
            "message": f"{inserted} paciente{'s' if inserted != 1 else ''} cargado{'s' if inserted != 1 else ''} con exito. {len(all_errors)} error{'es' if len(all_errors) != 1 else ''} encontrado{'s' if len(all_errors) != 1 else ''}.",
        })

    except Exception as e:
        db.session.rollback()
        return jsonify({"error": f"Error procesando el archivo: {e}"}), 500

    finally:
        # Limpiar archivo temporal
        try:
            os.remove(filepath)
        except OSError:
            pass


@app.route("/api/download-template/<tipo>")
def download_template(tipo):
    """Descarga una plantilla de ejemplo (Excel o Word)."""
    if tipo == "excel":
        path = os.path.join(TEMPLATE_FOLDER_DL, "plantilla_pacientes.xlsx")
        mimetype = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        filename = "plantilla_pacientes.xlsx"
    elif tipo == "word":
        path = os.path.join(TEMPLATE_FOLDER_DL, "plantilla_pacientes.docx")
        mimetype = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        filename = "plantilla_pacientes.docx"
    else:
        return jsonify({"error": "Tipo de plantilla no valido."}), 400

    if not os.path.exists(path):
        return jsonify({"error": "La plantilla no esta disponible."}), 404

    return send_file(path, mimetype=mimetype, as_attachment=True, download_name=filename)


# ---------------------------------------------------------------------------
# Rutas: Panel de Superusuario (Admin)
# ---------------------------------------------------------------------------
@app.route("/admin/dashboard")
@login_required
def admin_dashboard():
    """Panel de control del superusuario para gestionar personal de salud."""
    if current_user.role != "admin":
        flash("Acceso denegado. Se requieren permisos de administrador.", "error")
        return redirect(url_for("index"))
        
    staff_users = User.query.filter_by(role="health_staff").order_by(User.fecha_registro.desc()).all()
    unapproved_centers = HealthCenter.query.filter_by(is_approved=False).all()
    approved_centers = HealthCenter.query.filter_by(is_approved=True).order_by(HealthCenter.nombre.asc()).all()
    
    return render_template("admin_dashboard.html", users=staff_users, unapproved_centers=unapproved_centers, approved_centers=approved_centers)


@app.route("/admin/action/<action>/<int:user_id>", methods=["POST"])
@login_required
def admin_action(action, user_id):
    """Acciones CRUD sobre el personal de salud."""
    if current_user.role != "admin":
        return jsonify({"error": "Acceso denegado"}), 403
        
    user = User.query.get(user_id)
    if not user:
        return jsonify({"error": "Usuario no encontrado"}), 404
        
    if user.id == current_user.id:
        return jsonify({"error": "No puede modificar su propia cuenta desde aquí"}), 400
        
    if action == "approve":
        user.account_status = "approved"
        db.session.commit()
        return jsonify({"success": True, "message": "Usuario aprobado exitosamente."})
        
    elif action == "restrict":
        user.account_status = "restricted"
        db.session.commit()
        return jsonify({"success": True, "message": "Acceso restringido al usuario."})
        
    elif action == "delete":
        try:
            # Eliminar archivo de carnet si existe
            if user.id_card_path:
                filepath = os.path.join(app.config["ID_UPLOAD_FOLDER"], user.id_card_path)
                if os.path.exists(filepath):
                    os.remove(filepath)
            
            db.session.delete(user)
            db.session.commit()
            return jsonify({"success": True, "message": "Usuario eliminado completamente."})
        except Exception as e:
            db.session.rollback()
            return jsonify({"error": str(e)}), 500
            
    elif action == "edit":
        nombre = request.form.get("nombre", "").strip()
        centro = request.form.get("centro_medico_asociado", "").strip()
        if nombre:
            if not re.match(r"^[a-zA-ZáéíóúÁÉÍÓÚñÑ\s]{2,20}$", nombre):
                return jsonify({"error": "El nombre debe contener únicamente letras y tener entre 2 y 20 caracteres."}), 400
            user.nombre = nombre
        if centro:
            if not re.match(r"^[a-zA-ZáéíóúÁÉÍÓÚñÑ\s]{8,50}$", centro):
                return jsonify({"error": "El centro de salud debe contener únicamente letras y tener entre 8 y 50 caracteres."}), 400
            user.centro_medico_asociado = centro
        db.session.commit()
        return jsonify({"success": True, "message": "Datos actualizados."})
        
    return jsonify({"error": "Acción no válida"}), 400


@app.route("/admin/health_center/action/<action>/<int:center_id>", methods=["POST"])
@login_required
def admin_health_center_action(action, center_id):
    """Acciones CRUD sobre solicitudes de centros de salud."""
    if current_user.role != "admin":
        return jsonify({"error": "Acceso denegado"}), 403
        
    center = HealthCenter.query.get(center_id)
    if not center:
        return jsonify({"error": "Centro de salud no encontrado"}), 404
        
    if action == "approve":
        latitud_str = request.form.get("latitud", "").strip()
        longitud_str = request.form.get("longitud", "").strip()
        nombre = request.form.get("nombre", "").strip()
        tipo = request.form.get("tipo", "").strip()
        
        if nombre: center.nombre = nombre
        if tipo: center.tipo = tipo
        if latitud_str: center.latitud = float(latitud_str)
        if longitud_str: center.longitud = float(longitud_str)
        
        center.is_approved = True
        db.session.commit()
        return jsonify({"success": True, "message": "Centro de salud aprobado."})
        
    elif action == "reject":
        # Check if they want to merge
        merge_to_id_str = request.form.get("merge_to_id", "").strip()
        if merge_to_id_str:
            merge_to_id = int(merge_to_id_str)
            target_center = HealthCenter.query.get(merge_to_id)
            if target_center:
                # Reasign patients
                for reg in center.registros:
                    reg.health_center_id = target_center.id
                db.session.commit()
        
        db.session.delete(center)
        db.session.commit()
        return jsonify({"success": True, "message": "Solicitud rechazada/fusionada y eliminada."})

    return jsonify({"error": "Acción no válida"}), 400


@app.route("/admin/health_center/add", methods=["POST"])
@login_required
def admin_health_center_add():
    if current_user.role != "admin":
        return jsonify({"error": "No autorizado"}), 403

    nombre = request.form.get("nombre", "").strip()
    tipo = request.form.get("tipo", "").strip()
    latitud_str = request.form.get("latitud", "").strip()
    longitud_str = request.form.get("longitud", "").strip()

    if not nombre or not tipo:
        return jsonify({"error": "Nombre y tipo son obligatorios"}), 400

    # Sanitizar y validar Nombre de Centro: min 8 max 50, letras únicamente
    if not re.match(r"^[a-zA-ZáéíóúÁÉÍÓÚñÑ\s]{8,50}$", nombre):
        return jsonify({"error": "El nombre del centro de salud debe contener únicamente letras y tener entre 8 y 50 caracteres."}), 400

    latitud = float(latitud_str) if latitud_str else None
    longitud = float(longitud_str) if longitud_str else None

    nuevo_centro = HealthCenter(
        nombre=nombre,
        tipo=tipo,
        latitud=latitud,
        longitud=longitud,
        is_approved=True  # Automatically approved since an admin is creating it
    )

    db.session.add(nuevo_centro)
    db.session.commit()

    return jsonify({"success": True, "message": "Centro agregado exitosamente."})


@app.route("/admin/view_id/<filename>")
@login_required
def view_id(filename):
    """Permite al administrador ver el carnet de salud."""
    if current_user.role != "admin":
        flash("Acceso denegado.", "error")
        return redirect(url_for("index"))
        
    return send_file(os.path.join(app.config["ID_UPLOAD_FOLDER"], filename))


# ---------------------------------------------------------------------------
# Arranque
# ---------------------------------------------------------------------------
if __name__ == "__main__":
    with app.app_context():
        db.create_all()
        generate_templates()
        print("[OK] Base de datos inicializada.")
        print("[OK] Plantillas de descarga generadas.")
    app.run(debug=True, host="0.0.0.0", port=5000)
